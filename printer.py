import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches 

def create_pie(data):
    print('sedang buat pie chart')
    # Labels 
    labels = []
    nums = []
    for d in data:
        labels.append(d['namaGangguan'])
        nums.append(d['jumlahGangguan'])
    fig1, ax1 = plt.subplots()
    ax1.pie(nums, labels=labels, autopct='%1.1f%%',
            shadow=False, startangle=90)
    ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
    # plt.show()
    plt.savefig('pie.png')
    plt.close()

def create_line(data):
    print('sedang buat line chart')
    plt.xlabel('Tanggal')
    plt.ylabel('Jumlah Berita')
    #plt.title('Judul')
    date = data['axisx']

    for d in data['result']:
        jumlah = d['jumlahPerInterval']
        plt.plot(date, jumlah, label=d['namaGangguan'])
    plt.legend(loc='upper left', frameon=False)
    plt.xticks(rotation=-22)
    plt.savefig('line.png')
    # plt.show()
    plt.close()

def create_hist(data):
    date = [''] + data['axisx']
    rows = []
    for d in data['result']:
        rname = d['namaGangguan']
        rows.append([rname] + d['jumlahPerInterval'])
    plt.table(cellText=rows,colLabels=date,loc='center')
    #plt.title('Tabel Hasil Analisis')
    plt.axis('off')
    plt.axis('tight')
    plt.savefig('table.png')
    # plt.show()
    plt.close()

def createLaporan(pie_data, line_data, start, end):
    create_pie(pie_data)
    create_line(line_data)
    create_hist(line_data)

    start = line_data['axisx'][0]
    end = line_data['axisx'][len(line_data['axisx'])-1]
    # buat docx
    doc = Document()

    # kasih heading 
    doc.add_heading('Ringkasan Hasil Analisis \n{} s/d {}'.format(start, end), 0)

    # memasukkan chart ke docx
    doc.add_heading('Pie chart persebaran gangguan', level=1)
    doc.add_picture('pie.png', width=Inches(6))
    doc.add_heading('Line chart persebaran gangguan', level=1)
    doc.add_picture('line.png', width=Inches(6))
    doc.add_heading('Tabel persebaran gangguan', level=1)
    doc.add_picture('table.png', width=Inches(6))

    # menyimpan docx 
    doc.save('lapor.docx')
